from rest_framework import viewsets, permissions, decorators, status, mixins
from rest_framework.response import Response
from django.http import FileResponse, HttpResponse
from rest_framework.decorators import action
from docx import Document
from io import BytesIO
from drf_spectacular.utils import (
    extend_schema, extend_schema_view, OpenApiResponse, OpenApiTypes
)
from rest_framework.parsers import MultiPartParser, FormParser

from coins.models import Coin
from tags.models import Tag
from favorites.models import Favorite
from users.models import User
from .serializers import (
    TagSerializer,
    CoinReadSerializer, CoinWriteSerializer,
    FavoriteSerializer,
)
from .filters import CoinFilter
from .exceptions import PermissionDeniedError


# ---------------------------------------------------
# Теги
# ---------------------------------------------------

class TagViewSet(viewsets.ModelViewSet):
    queryset = Tag.objects.all()
    serializer_class = TagSerializer

    def get_permissions(self):
        if self.request.method in permissions.SAFE_METHODS:
            return [permissions.AllowAny()]
        return [permissions.IsAdminUser()]


# ---------------------------------------------------
# Монеты
# ---------------------------------------------------

@extend_schema_view(
    list=extend_schema(
        responses=CoinReadSerializer(many=True),
        description="Список всех монет"
    ),
    retrieve=extend_schema(
        responses=CoinReadSerializer,
        description="Получить детали монеты"
    ),
    destroy=extend_schema(
        responses={204: OpenApiTypes.NONE},
        description="Удалить монету"
    ),
)
class CoinViewSet(viewsets.ModelViewSet):
    queryset = Coin.objects.select_related('author').prefetch_related('tags')
    permission_classes = (permissions.IsAuthenticatedOrReadOnly,)
    filterset_class = CoinFilter
    parser_classes = (MultiPartParser, FormParser)
    search_fields = ('name', 'description')
    ordering_fields = ('estimated_value', 'pub_date')

    @extend_schema(
        request=CoinWriteSerializer,
        responses={201: CoinReadSerializer},
        description="Создать монету с загрузкой изображения"
    )
    def create(self, request, *args, **kwargs):
        return super().create(request, *args, **kwargs)

    @extend_schema(
        request=CoinWriteSerializer,
        responses=CoinReadSerializer,
        description="Полностью обновить монету (PUT) с загрузкой изображения"
    )
    def update(self, request, *args, **kwargs):
        return super().update(request, *args, **kwargs)

    @extend_schema(
        request=CoinWriteSerializer,
        responses=CoinReadSerializer,
        description="Частично обновить монету (PATCH) с загрузкой изображения"
    )
    def partial_update(self, request, *args, **kwargs):
        return super().partial_update(request, *args, **kwargs)

    def get_serializer_class(self):
        if self.action in ('create', 'update', 'partial_update'):
            return CoinWriteSerializer
        return CoinReadSerializer

    def perform_create(self, serializer):
        serializer.save(author=self.request.user)

    def perform_update(self, serializer):
        coin = self.get_object()
        user = self.request.user
        if not (user.is_staff or coin.author == user):
            raise PermissionDeniedError("Редактировать можно лишь свои монеты")
        serializer.save()

    def perform_destroy(self, instance):
        user = self.request.user
        if not (user.is_staff or instance.author == user):
            raise PermissionDeniedError("Удалять можно лишь свои монеты")
        instance.delete()


# ---------------------------------------------------
# Избранное
# ---------------------------------------------------

@extend_schema_view(
    list=extend_schema(
        responses=CoinReadSerializer(many=True),
        description="Список монет в избранном текущего пользователя"
    ),
    create=extend_schema(
        request=FavoriteSerializer,
        responses={201: FavoriteSerializer},
        description="Добавить монету в избранное"
    ),
    destroy=extend_schema(
        responses={204: OpenApiTypes.NONE},
        description="Удалить монету из избранного"
    ),
    download=extend_schema(
        request=None,
        responses=OpenApiResponse(
            response=OpenApiTypes.BINARY,
            description="Word-файл со списком избранного"
        ),
    ),
)
class FavoriteViewSet(
    mixins.CreateModelMixin,
    mixins.ListModelMixin,
    viewsets.GenericViewSet,
):
    queryset = Favorite.objects.all()
    serializer_class = FavoriteSerializer
    permission_classes = (permissions.IsAuthenticated,)

    def get_queryset(self):
        # swagger_fake_view нужен, чтобы spectacular не вывалился на анониме
        if getattr(self, "swagger_fake_view", False):
            return Favorite.objects.none()
        return Favorite.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    @action(detail=False, methods=['delete'], url_path='by_coin')

    def delete_by_coin(self, request):
        coin_id = request.query_params.get('coin')
        if not coin_id:
            return Response(
                {'detail': 'coin обязателен как query param'},
                status=status.HTTP_400_BAD_REQUEST
            )
        instance = Favorite.objects.filter(
            user=request.user,
            coin_id=coin_id
        ).first()
        if not instance:
            return Response({'detail': 'Избранное не найдено'}, status=status.HTTP_404_NOT_FOUND)
        instance.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)
    def list(self, request, *args, **kwargs):
        coins = Coin.objects.filter(in_favorites__user=request.user)
        return Response(CoinReadSerializer(coins, many=True).data)

    @decorators.action(detail=False, methods=['get'])
    def download(self, request):
        coins = Coin.objects.filter(in_favorites__user=request.user)
        doc = Document()
        doc.add_heading('Список избранного', level=1)
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = '№'
        hdr[1].text = 'Название'
        hdr[2].text = 'Теги'
        hdr[3].text = 'Цена, ₽'
        hdr[4].text = 'Продавец'

        for i, coin in enumerate(coins, start=1):
            cells = table.add_row().cells
            cells[0].text = str(i)
            cells[1].text = coin.name
            cells[2].text = ', '.join(t.name for t in coin.tags.all())
            cells[3].text = str(coin.estimated_value)
            cells[4].text = coin.author.username

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        response = HttpResponse(
            buf.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="favorites.docx"'
        return response